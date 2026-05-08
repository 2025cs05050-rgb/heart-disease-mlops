"""Generate the MLOps Assignment-1 Word (.docx) report.

Mirrors the structure of ``generate_report.py`` (PDF) but emits a
``.docx`` so the deliverable matches the assignment brief, which asks
for a ``doc/docx`` file. The prose is reused verbatim from
``report_content.py`` via a shadow-helpers pattern: ``section_*``
functions emit lightweight tuples that this module renders into native
Word constructs (paragraphs, runs, tables, images).

Run:
    python scripts/generate_docx.py
Output:
    reports/MLOps_Assignment1_Report.docx
"""
from __future__ import annotations

import json
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from report_content import (  # noqa: E402
    cover_block,
    section_appendix,
    section_architecture,
    section_ci_cd,
    section_conclusion,
    section_containerization,
    section_deliverables,
    section_eda,
    section_executive_summary,
    section_kubernetes,
    section_mlflow,
    section_modeling,
    section_monitoring,
    section_pipeline,
    section_problem,
    section_production_deployment_intro,
    section_repo,
)

ROOT = Path(__file__).resolve().parents[1]
METRICS = json.loads((ROOT / "reports" / "metrics.json").read_text())
FIGS = ROOT / "reports" / "figures"
SHOTS = ROOT / "screenshots"
OUT = ROOT / "reports" / "MLOps_Assignment1_Report.docx"

PRIMARY = RGBColor(0x1F, 0x4E, 0x79)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GREY = RGBColor(0x55, 0x55, 0x55)
MONO_FONT = "Consolas"
BODY_FONT = "Calibri"

# --------------------------------------------------------------------- markup
TAG_RE = re.compile(r"<(/?)(b|i|code|br\s*/?|font[^>]*)>", re.IGNORECASE)


def _decode(text: str) -> str:
    return (text.replace("&nbsp;", "\u00A0")
                .replace("&lt;", "<").replace("&gt;", ">").replace("&amp;", "&"))


def parse_inline(text: str) -> list[tuple[str, dict]]:
    """Tokenise the ReportLab-style markup into ``(text, attrs)`` runs.

    ``\\n`` inside a chunk represents a line break that the renderer
    must convert into a Word ``<w:br/>``.
    """
    text = _decode(text)
    runs: list[tuple[str, dict]] = []
    stack: list[dict] = []
    pos = 0

    def attrs_now() -> dict:
        merged: dict = {}
        for s in stack:
            merged.update(s)
        return merged

    while True:
        m = TAG_RE.search(text, pos)
        if not m:
            if pos < len(text):
                runs.append((text[pos:], attrs_now()))
            break
        if m.start() > pos:
            runs.append((text[pos:m.start()], attrs_now()))
        closing = m.group(1) == "/"
        tag = m.group(2).lower()
        if tag.startswith("br"):
            runs.append(("\n", attrs_now()))
        elif closing:
            if stack:
                stack.pop()
        else:
            if tag == "b":
                stack.append({"bold": True})
            elif tag == "i":
                stack.append({"italic": True})
            elif tag == "code":
                stack.append({"mono": True})
            elif tag.startswith("font"):
                stack.append({"mono": True})
        pos = m.end()
    return runs


def add_runs(paragraph, text: str, base: dict | None = None) -> None:
    """Render parsed inline markup into a docx paragraph."""
    base = base or {}
    for chunk, attrs in parse_inline(text):
        merged = {**base, **attrs}
        if not chunk:
            continue
        # Split on \n so each linebreak produces a w:br element.
        parts = chunk.split("\n")
        for i, part in enumerate(parts):
            if part:
                run = paragraph.add_run(part)
                if merged.get("bold"):
                    run.bold = True
                if merged.get("italic"):
                    run.italic = True
                if merged.get("mono"):
                    run.font.name = MONO_FONT
                    run.font.size = Pt(9)
                else:
                    run.font.name = BODY_FONT
                if "color" in merged:
                    run.font.color.rgb = merged["color"]
                if "size" in merged:
                    run.font.size = merged["size"]
            if i < len(parts) - 1:
                paragraph.add_run().add_break()


# --------------------------------------------------------------------- styles
STYLE_SPECS = {
    "title":    {"size": Pt(24), "color": PRIMARY, "bold": True,
                 "align": WD_ALIGN_PARAGRAPH.CENTER, "space_after": Pt(14)},
    "subtitle": {"size": Pt(12), "color": GREY,
                 "align": WD_ALIGN_PARAGRAPH.CENTER, "space_after": Pt(4)},
    "h1":       {"size": Pt(15), "color": PRIMARY, "bold": True,
                 "space_before": Pt(12), "space_after": Pt(6)},
    "h2":       {"size": Pt(11.5), "color": ACCENT, "bold": True,
                 "space_before": Pt(8), "space_after": Pt(3)},
    "body":     {"size": Pt(10.5), "align": WD_ALIGN_PARAGRAPH.JUSTIFY,
                 "space_after": Pt(4)},
    "bullet":   {"size": Pt(10.5), "space_after": Pt(2),
                 "left_indent": Cm(0.6)},
    "caption":  {"size": Pt(9), "color": GREY, "italic": True,
                 "align": WD_ALIGN_PARAGRAPH.CENTER, "space_after": Pt(8)},
    "code":     {"size": Pt(9), "mono": True, "space_after": Pt(6)},
}


def _apply_paragraph_format(p, spec: dict) -> None:
    if "align" in spec:
        p.alignment = spec["align"]
    fmt = p.paragraph_format
    if "space_before" in spec:
        fmt.space_before = spec["space_before"]
    if "space_after" in spec:
        fmt.space_after = spec["space_after"]
    if "left_indent" in spec:
        fmt.left_indent = spec["left_indent"]


def add_paragraph(doc, text: str, style_key: str) -> None:
    spec = STYLE_SPECS.get(style_key, STYLE_SPECS["body"])
    p = doc.add_paragraph()
    _apply_paragraph_format(p, spec)
    if style_key == "bullet":
        if text.startswith("• "):
            text = text[2:]
        _set_bullet(p)
    base = {k: v for k, v in spec.items()
            if k in ("size", "color", "bold", "italic", "mono")}
    add_runs(p, text, base=base)


def _set_bullet(p) -> None:
    """Apply a list-paragraph bullet via direct numPr XML."""
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "1")
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


# --------------------------------------------------------------------- table
def add_metrics_table(doc, metrics: dict) -> None:
    by_name = {r["name"]: r for r in metrics["all_results"]}
    keys = ("test_accuracy", "test_precision", "test_recall",
            "test_f1", "test_roc_auc")
    table = doc.add_table(rows=1 + len(keys), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(("Metric", "Logistic Regression", "Random Forest")):
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(hdr[i], "1F4E79")
    for r, key in enumerate(keys, start=1):
        cells = table.rows[r].cells
        cells[0].text = key.replace("test_", "").upper()
        cells[1].text = (
            f"{by_name['logistic_regression']['test_metrics'][key]:.4f}")
        cells[2].text = (
            f"{by_name['random_forest']['test_metrics'][key]:.4f}")
        for c in cells:
            for para in c.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


# --------------------------------------------------------------------- image
def add_image(doc, path: Path, caption: str, width_cm: float = 14.0) -> None:
    if not path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"[Figure missing: {path.name}]")
        run.italic = True
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    _apply_paragraph_format(cap, STYLE_SPECS["caption"])
    add_runs(cap, caption, base={"italic": True, "size": Pt(9),
                                 "color": GREY})


# --------------------------------------------------------------------- shadow helpers
SHADOW_STYLES = {k: k for k in STYLE_SPECS}


def make_helpers() -> dict:
    return {
        "Paragraph": lambda text, style: ("p", text, style),
        "Spacer": lambda w, h: ("space", h),
        "today": date.today().isoformat(),
    }


def render(doc, flowables) -> None:
    for f in flowables:
        if not isinstance(f, tuple):
            continue
        kind = f[0]
        if kind == "p":
            add_paragraph(doc, f[1], f[2])
        elif kind == "space" and f[1] >= 20:
            doc.add_paragraph()


# --------------------------------------------------------------------- numbering
_NUMBERING_XML = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:abstractNum w:abstractNumId="0">
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="bullet"/>
      <w:lvlText w:val="\u2022"/>
      <w:lvlJc w:val="left"/>
      <w:pPr><w:ind w:left="360" w:hanging="360"/></w:pPr>
      <w:rPr><w:rFonts w:ascii="Symbol" w:hAnsi="Symbol"/></w:rPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1"><w:abstractNumId w:val="0"/></w:num>
</w:numbering>
"""


def _ensure_numbering(doc) -> None:
    """Attach a minimal numbering.xml part so bullet lists render."""
    from docx.opc.constants import CONTENT_TYPE, RELATIONSHIP_TYPE
    from docx.opc.packuri import PackURI
    from docx.opc.part import Part

    part = doc.part
    if any(rel.reltype == RELATIONSHIP_TYPE.NUMBERING
           for rel in part.rels.values()):
        return
    numbering_part = Part(
        PackURI("/word/numbering.xml"),
        CONTENT_TYPE.WML_NUMBERING,
        _NUMBERING_XML.encode("utf-8"),
        part.package,
    )
    part.relate_to(numbering_part, RELATIONSHIP_TYPE.NUMBERING)


# --------------------------------------------------------------------- footer
def _setup_footer(doc) -> None:
    section = doc.sections[0]
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "MLOps Assignment-1  —  Heart Disease Risk Pipeline  —  Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    page_run = p.add_run()
    page_run.font.size = Pt(8)
    page_run.font.color.rgb = GREY
    page_run._r.append(fld_begin)
    page_run._r.append(instr)
    page_run._r.append(fld_end)


def _setup_page(doc) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)


# --------------------------------------------------------------------- build
def build() -> Path:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    _setup_page(doc)
    _ensure_numbering(doc)
    _setup_footer(doc)

    helpers = make_helpers()
    styles = SHADOW_STYLES

    render(doc, cover_block(styles, helpers))
    doc.add_page_break()

    render(doc, section_executive_summary(styles, helpers, METRICS))
    render(doc, section_problem(styles, helpers))
    doc.add_page_break()

    render(doc, section_repo(styles, helpers))
    render(doc, section_eda(styles, helpers))
    doc.add_page_break()

    render(doc, section_pipeline(styles, helpers))
    render(doc, section_modeling(styles, helpers, METRICS, lambda: []))
    add_metrics_table(doc, METRICS)
    auc = METRICS["all_results"][1]["test_metrics"]["test_roc_auc"]
    add_image(doc, FIGS / "cm_random_forest.png",
              "Figure 1 — Confusion matrix on the held-out test set "
              "(Random Forest).", width_cm=10)
    add_image(doc, FIGS / "roc_random_forest.png",
              f"Figure 2 — ROC curve, test AUC = {auc:.4f}.", width_cm=10)
    doc.add_page_break()

    render(doc, section_mlflow(styles, helpers))
    add_image(doc, FIGS / "screenshot_mlflow.png",
              "Figure 3 — MLflow run-comparison view for the heart-disease "
              "experiment (rendered from logged metrics).", width_cm=15)
    render(doc, section_ci_cd(styles, helpers))
    doc.add_page_break()

    render(doc, section_containerization(styles, helpers))
    add_image(doc, FIGS / "screenshot_swagger.png",
              "Figure 4 — Swagger UI at /docs showing the five service "
              "endpoints.", width_cm=15)
    render(doc, section_kubernetes(styles, helpers))
    doc.add_page_break()

    render(doc, section_monitoring(styles, helpers))
    add_image(doc, FIGS / "screenshot_grafana.png",
              "Figure 5 — Pre-provisioned Grafana dashboard ML › Heart "
              "Disease API: stat row, request rate, latency percentiles, "
              "class balance and HTTP status codes.", width_cm=16)
    doc.add_page_break()

    render(doc, section_production_deployment_intro(styles, helpers))
    shots = [
        ("01_cluster_context.png",
         "Figure 6 — kubectl get nodes showing the two-node cluster "
         "(controlplane + node01) both Ready."),
        ("02_docker_image.png",
         "Figure 7 — docker images heart-disease-api: image produced "
         "by the multi-stage build."),
        ("03_pods_running.png",
         "Figure 8 — kubectl -n heart-disease get pods -o wide: both "
         "replicas 1/1 Running across the two nodes."),
        ("04_deployment_describe.png",
         "Figure 9 — kubectl describe deployment: RollingUpdate "
         "strategy, 2/2 available, probes wired."),
        ("05_service_nodeport.png",
         "Figure 10 — kubectl get svc: NodePort service exposed for "
         "external traffic."),
        ("06_hpa.png",
         "Figure 11 — kubectl get hpa: active HPA (2-5 replicas, "
         "70 % CPU target) with live metrics."),
        ("07_curl_health_predict.png",
         "Figure 12 — Live curl calls: /health and /predict returning "
         "valid JSON."),
    ]
    for filename, caption in shots:
        add_image(doc, SHOTS / filename, caption, width_cm=15)
    doc.add_page_break()

    render(doc, section_architecture(styles, helpers))
    doc.add_page_break()

    render(doc, section_conclusion(styles, helpers, METRICS))
    render(doc, section_appendix(styles, helpers))
    render(doc, section_deliverables(styles, helpers))

    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    out = build()
    print(f"DOCX written -> {out}  ({out.stat().st_size / 1024:.1f} KB)")
